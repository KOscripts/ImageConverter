from docx import Document
from docx.shared import Inches
import os
from datetime import date


def convert(FilePath, FileName=''):

    today = date.today()
    currentDate = today.strftime('%m_%d_%y')
    path = fr'{FilePath}'

    images = os.listdir(path)

    document = Document()

    if len(images) != 0:
        try:
            for image in images:
                currentValue = 0

                im = os.path.splitext(images[currentValue])[0]

                document.add_paragraph(f'{im}')

                document.add_picture(fr'{path}\{image}',
                                     width=Inches(7), height=Inches(6))
                document.add_page_break()

                currentValue += 1

            document.save(f'{str(FileName)}_{currentDate}.docx')
        except:
            pass
    else:
        pass
